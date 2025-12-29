"""
========================
T2D2 SDK Client Library
========================
"""

import json
import logging
import os
# pylint: disable=wildcard-import, unused-wildcard-import
import random
import string
from collections import defaultdict
from datetime import datetime
from enum import Enum, auto
from urllib.parse import urlencode, urlparse

import boto3
import requests
import sentry_sdk

# Set up logger
logger = logging.getLogger(__name__)

# Configure logging if not already configured
if not logger.handlers:
    handler = logging.StreamHandler()
    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    logger.setLevel(logging.INFO)  # Use INFO level only - DEBUG messages will be filtered
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from PIL import Image as PILImage
from condition_report_service import ImageAnnotationCropper


TIMEOUT = 60
BASE_URL = os.getenv("T2D2_API_URL", "https://api-v3.t2d2.ai/api/")
# DEV https://api-v3-dev.t2d2.ai/api/


####################################################################################################
sentry_sdk.init(
    dsn="https://fdaf778d002a179fa64754611aa9ace0@sentry.t2d2.ai/5",
    # Set traces_sample_rate to 1.0 to capture 100%
    # of transactions for performance monitoring.
    traces_sample_rate=1.0,
)
####################################################################################################

####################################################################################################
# COMMON HELPER FUNCTIONS
####################################################################################################
def random_string(length: int = 6) -> str:
    """
    Generate a random string of fixed length.
    
    :param length: The length of the random string to generate
    :type length: int
    :default length: 6
    
    :return: A random string consisting of lowercase ASCII letters
    :rtype: str
    
    :example:
        >>> random_string()
        'abcdef'
        >>> random_string(10)
        'abcdefghij'
    """
        
    letters = string.ascii_lowercase
    return "".join(random.choice(letters) for i in range(length))


def random_color() -> str:
    """
    Generate a random hexadecimal color code.
    
    :return: A random color in hexadecimal format (e.g., '#FF5733')
    :rtype: str
    
    :example:
        >>> random_color()
        '#A2B9C7'
    """
    r = lambda: random.randint(0, 255)
    return "#%02X%02X%02X" % (r(), r(), r())


def ts2date(ts):
    """
    Convert a Unix timestamp to a human-readable date string.
    
    :param ts: Unix timestamp to convert
    :type ts: int or float
    
    :return: Formatted date string in 'YYYY-MM-DD HH:MM:SS' format
    :rtype: str
    
    :example:
        >>> ts2date(1609459200)
        '2021-01-01 00:00:00'
    """
    return datetime.fromtimestamp(ts).strftime("%Y-%m-%d %H:%M:%S")


def download_file(url: str, file_path: str):
    """
    Download a file from an S3 URL to a local path.
    
    This function parses the S3 URL to extract bucket and key information,
    then downloads the file to the specified local path.
    
    :param url: S3 URL of the file to download
    :type url: str
    :param file_path: Local path where the file should be saved
    :type file_path: str
    
    :return: Dictionary containing success status and message
    :rtype: dict
    
    :raises Exception: If there is an error during download
    
    :example:
        >>> download_file('s3://my-bucket/path/to/file.jpg', '/local/path/file.jpg')
        {'success': True, 'message': 'File downloaded'}
    """
    try:
        s3 = boto3.client("s3")
        parsed_url = urlparse(url)
        bucket = parsed_url.netloc.split(".")[0]
        key = parsed_url.path[1:]
        s3.download_file(bucket, key, file_path)
        return {"success": True, "message": "File downloaded"}
    except Exception as e:
        return {
            "success": False,
            "message": f"{str(e)} \n{bucket} \n{key} \n{file_path}",
        }


def upload_file(file_path: str, url: str):
    """
    Upload a file from a local path to an S3 URL.
    
    This function parses the S3 URL to extract bucket and key information,
    then uploads the file from the specified local path with public-read ACL.
    
    :param file_path: Local path of the file to upload
    :type file_path: str
    :param url: S3 URL where the file should be uploaded
    :type url: str
    
    :return: Dictionary containing success status and message
    :rtype: dict
    
    :raises Exception: If there is an error during upload
    
    :example:
        >>> upload_file('/local/path/file.jpg', 's3://my-bucket/path/to/file.jpg')
        {'success': True, 'message': 'File uploaded'}
    """
    try:
        s3 = boto3.client("s3")
        parsed_url = urlparse(url)
        bucket = parsed_url.netloc.split(".")[0]
        key = parsed_url.path[1:]
        s3.upload_file(file_path, bucket, key, ExtraArgs={"ACL": "public-read"})
        return {"success": True, "message": "File uploaded"}
    except Exception as e:
        return {
            "success": False,
            "message": f"{str(e)} \n{bucket} \n{key} \n{file_path}",
        }


####################################################################################################
# Helper function for adding hyperlinks to docx
####################################################################################################
def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph in a Word document.
    
    :param paragraph: The paragraph to add the hyperlink to
    :param url: The URL to link to
    :param text: The text to display for the hyperlink
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Style the hyperlink
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue color
    rPr.append(color)
    
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

####################################################################################################
class RequestType(Enum):
    """Request types"""

    GET = auto()
    PUT = auto()
    POST = auto()
    DELETE = auto()


####################################################################################################

####################################################################################################
class T2D2(object):
    """
    A class to interact with the T2D2 API for structural inspection data management.
    
    This class provides a comprehensive interface to the T2D2 API, enabling authentication,
    project management, and manipulation of various data types including images, annotations,
    drawings, videos, reports, and more.
    
    :ivar base_url: The base URL for the T2D2 API endpoints
    :vartype base_url: str
    :ivar headers: Headers sent with each API request (includes content type and authorization)
    :vartype headers: dict
    :ivar s3_base_url: The base URL for S3 storage where data is stored or retrieved
    :vartype s3_base_url: str
    :ivar aws_region: The AWS region for S3 interactions
    :vartype aws_region: str
    :ivar bucket: The name of the S3 bucket used for storing data
    :vartype bucket: str
    :ivar access_token: The token used for authenticating API requests
    :vartype access_token: str
    :ivar api_key: The API key used for authenticating with the T2D2 API
    :vartype api_key: str
    :ivar project: A dictionary representing the current project's data
    :vartype project: dict
    :ivar debug: A flag to enable or disable debug mode
    :vartype debug: bool
    """

    base_url: str
    headers: dict
    s3_base_url: str
    aws_region: str
    bucket: str
    access_token: str
    api_key: str
    project: dict = None
    debug: bool = True

    def __init__(self, credentials, base_url=BASE_URL):
        """
        Initialize the T2D2 client and authenticate with the API.
        
        Creates a new T2D2 client instance, configures it with the provided base URL,
        and authenticates using the supplied credentials. This method ensures the base URL 
        ends with a forward slash and sets up the default content type header.
        
        :param credentials: Authentication credentials dictionary which can include:
                        - 'access_token': Direct token for authentication
                        - 'username'/'password': For login-based authentication
                        - 'api_key': API key for authentication
        :type credentials: dict
        :param base_url: Base URL for T2D2 API endpoints
        :type base_url: str
        :default base_url: Uses the global BASE_URL value
        
        :raises ValueError: If the base_url does not end with a forward slash or if 
                        authentication fails
        
        :example:
            >>> client = T2D2({'api_key': 'your-api-key'})
            >>> client = T2D2({'access_token': 'your-token'})
            >>> client = T2D2({'username': 'user@example.com', 'password': 'pass123'})
        """
        if not base_url.endswith("/"):
            base_url += "/"
        self.base_url = base_url
        self.headers = {"Content-Type": "application/json"}
        logger.info(f"Initializing T2D2 client with base_url: {base_url}")
        self.login(credentials)
        self.project = {}
        logger.debug("T2D2 client initialized successfully")

    def request(
        self,
        url_suffix: str,
        req_type: RequestType = RequestType.GET,
        params=None,
        headers=None,
        data=None,
    ) -> dict:
        """Send a request and handle response"""
        url = self.base_url + url_suffix
        logger.debug(f"Making {req_type.name} request to: {url}")
        if headers is None:
            headers = {}
        if params is None:
            params = {}
        if data is None:
            data = {}

        headers.update(self.headers)
        params_enc = {}
        for key, val in params.items():
            if isinstance(val, list):
                params_enc[key] = json.dumps(val)
            else:
                params_enc[key] = val
        if req_type == RequestType.GET:
            res = requests.get(
                url,
                headers=headers,
                params=urlencode(params_enc),
                timeout=TIMEOUT,
            )
        elif req_type == RequestType.POST:
            res = requests.post(
                url,
                headers=headers,
                params=urlencode(params_enc),
                json=data,
                timeout=TIMEOUT,
            )
        elif req_type == RequestType.PUT:
            res = requests.put(
                url,
                headers=headers,
                params=urlencode(params_enc),
                json=data,
                timeout=TIMEOUT,
            )
        elif req_type == RequestType.DELETE:
            res = requests.delete(
                url,
                headers=headers,
                params=urlencode(params_enc),
                json=data,
                timeout=TIMEOUT,
            )
        else:
            raise ValueError("Request type not yet supported. Coming soon")

        if res.status_code in (200, 201):
            try:
                logger.debug(f"Request successful: {req_type.name} {url} - Status: {res.status_code}")
                return res.json()
            except Exception as e:
                logger.error(f"JSON Conversion Error for {url}: {e}")
                return {"content": res.content}
        else:
            logger.error(f"Request failed: {req_type.name} {url} - Status: {res.status_code}")
            if self.debug:
                logger.debug(f"URL: {req_type} {res.url}")
                logger.debug(f"HEADERS: {headers}")
                logger.debug(f"PARAMS: {params}")
                logger.debug(f"DATA: {data}")
                logger.debug(f"Response: {res.status_code} {res.content}")
            raise ValueError(f"Error code received: {res.status_code}")

    def login(self, credentials):
        """
        Authenticate with the T2D2 API using provided credentials.
        
        This method supports three authentication methods:
        1. Direct access token authentication
        2. Username/password login
        3. API key authentication
        
        Upon successful authentication, the appropriate authorization header 
        is set for future requests.
        
        :param credentials: Authentication credentials dictionary which must include one of:
                        - 'access_token': Direct token for authentication
                        - 'username' and 'password': For login-based authentication
                        - 'api_key': API key for authentication
        :type credentials: dict
        
        :return: None
        
        :raises ValueError: If the provided credentials are invalid, missing required fields,
                        or if the authentication request fails
        
        :example:
            >>> client.login({'access_token': 'your-access-token'})
            >>> client.login({'username': 'user@example.com', 'password': 'password123'})
            >>> client.login({'api_key': 'your-api-key'})
        """

        if "access_token" in credentials:
            # Directly use token
            logger.info("Authenticating with access token")
            self.access_token = credentials["access_token"]
            self.headers["Authorization"] = f"Bearer {self.access_token}"
            logger.debug("Access token authentication successful")

        elif "password" in credentials:
            # Login
            logger.info(f"Authenticating with username/password for user: {credentials.get('username', 'N/A')}")
            url = "auth/login"
            json_data = self.request(url, RequestType.POST, data=credentials)
            self.access_token = json_data["data"]["firebaseDetail"]["access_token"]
            self.headers["Authorization"] = f"Bearer {self.access_token}"
            logger.info("Username/password authentication successful")

        elif "api_key" in credentials:
            logger.info("Authenticating with API key")
            self.api_key = credentials["api_key"]
            self.headers["x-api-key"] = self.api_key
            logger.debug("API key authentication successful")
        else:
            logger.error("No valid credentials provided")
            raise ValueError("No valid credentials provided. Must include 'access_token', 'password'/'username', or 'api_key'")

        return

    def notify_user(self, title, message):
        """Send a notification to the user

        Args:
            title (str): The title of the notification.
            message (str): The message content of the notification.

        Returns:
            dict: A dictionary containing the response from the API.

        Raises:
            ValueError: If the notification could not be sent.

        """
        url = "notifications"
        payload = {"title": title, "message": message}
        return self.request(url, RequestType.POST, data=payload)

    ################################################################################################
    # Project Get/Set
    ################################################################################################
    def get_project(self, project_id=None):
        """
        Retrieve a project by ID or get a list of all projects.
        
        If project_id is provided, returns data for the specific project.
        If project_id is None, returns a list of all available projects.
        
        :param project_id: The ID of the specific project to retrieve
        :type project_id: str or None
        :default project_id: None
        
        :return: Project data dictionary for a specific project, or a list of projects
        :rtype: dict
        
        :raises ValueError: If the API request fails or the project_id is invalid
        
        :example:
            >>> # Get a specific project
            >>> project = client.get_project("project_123abc")
            >>> print(project["profile"]["name"])
            'Bridge Inspection'
            
            >>> # Get all projects
            >>> all_projects = client.get_project()
            >>> print(len(all_projects["project_list"]))
            5
        """
        if project_id is None:
            logger.info("Fetching all projects")
            url = "project"
        else:
            logger.info(f"Fetching project: {project_id}")
            url = f"project/{project_id}"
        json_data = self.request(url, RequestType.GET)
        logger.debug(f"Successfully retrieved project data")
        return json_data["data"]
    
    def create_project(self, name, location=None, address=None, latitude=None, 
                      longitude=None, country=None, country_code=None):
        """
        Create a new project in T2D2.
        
        This method creates a new project with the specified name and location information.
        You can either provide a complete location dictionary or individual location parameters.
        
        :param name: The name of the project to create
        :type name: str
        
        :param location: A complete location dictionary with address, latitude, longitude, 
                        country, and country_code. If provided, individual location parameters are ignored.
        :type location: dict or None
        :default location: None
        
        :param address: The address/location of the project
        :type address: str or None
        :default address: None
        
        :param latitude: The latitude coordinate of the project location
        :type latitude: float or None
        :default latitude: None
        
        :param longitude: The longitude coordinate of the project location
        :type longitude: float or None
        :default longitude: None
        
        :param country: The country where the project is located
        :type country: str or None
        :default country: None
        
        :param country_code: The country code (e.g., "UK", "US")
        :type country_code: str or None
        :default country_code: None
        
        :return: A dictionary containing the created project details
        :rtype: dict
        
        :raises ValueError: If the API request fails or required parameters are missing
        
        :example:
            >>> # Create project with individual location parameters
            >>> project = client.create_project(
            ...     name="120 Broadway",
            ...     address="Ahmedabad",
            ...     latitude=1234.89,
            ...     longitude=456.9,
            ...     country="United Kingdom",
            ...     country_code="UK"
            ... )
            >>> print(f"Created project: {project['profile']['name']} with ID: {project['id']}")
            
            >>> # Create project with location dictionary
            >>> location = {
            ...     "address": "Ahmedabad",
            ...     "latitude": 1234.89,
            ...     "longitude": 456.9,
            ...     "country": "United Kingdom",
            ...     "country_code": "UK"
            ... }
            >>> project = client.create_project(name="120 Broadway", location=location)
            >>> print(f"Created project ID: {project['id']}")
        """
        # Build location dictionary
        if location is None:
            location = {}
            if address is not None:
                location["address"] = address
            if latitude is not None:
                location["latitude"] = latitude
            if longitude is not None:
                location["longitude"] = longitude
            if country is not None:
                location["country"] = country
            if country_code is not None:
                location["country_code"] = country_code
        
        # Build payload
        payload = {
            "name": name,
            "location": location
        }
        
        logger.info(f"Creating new project: {name}")
        url = "project"
        json_data = self.request(url, RequestType.POST, data=payload)
        project_id = json_data["data"].get("id")
        logger.info(f"Project created successfully with ID: {project_id}")
        return json_data["data"]

    def set_project(self, project_id):
        """
        Set the current active project for the T2D2 client.
        
        This method retrieves the project data by ID, sets it as the active project,
        and configures related S3 storage settings for the client instance.
        
        :param project_id: The ID of the project to set as active
        :type project_id: str
        
        :return: None
        
        :raises ValueError: If the project_id is invalid or the project cannot be accessed
        
        :example:
            >>> client.set_project("project_123abc")
            >>> print(client.project["profile"]["name"])
            'Bridge Inspection'
        
        :note: After setting a project, S3 storage information is automatically configured
            for use with other methods that require file uploads or downloads.
        """
        logger.info(f"Setting project: {project_id}")
        json_data = self.request(f"project/{project_id}", RequestType.GET)
        if not json_data["success"]:
            logger.error(f"Failed to set project {project_id}: {json_data.get('message', 'Unknown error')}")
            raise ValueError(json_data["message"])

        project = json_data["data"]
        self.project = project
        logger.debug(f"Project set: {project.get('profile', {}).get('name', 'N/A')}")

        self.s3_base_url = project["config"]["s3_base_url"]
        self.aws_region = project["config"]["aws_region"]
        res = urlparse(self.s3_base_url)
        self.bucket = res.netloc.split(".")[0]
        logger.debug(f"S3 configuration set: bucket={self.bucket}, region={self.aws_region}")
        return

    def get_project_info(self):
        """
        Retrieve a summary of the current active project's information.
        
        This method returns a dictionary containing the essential information
        about the currently selected project, including its ID, name, address,
        description, creation details, and statistics.
        
        :return: A dictionary containing the project information with the following keys:
                - id: The project's unique identifier
                - name: The project's display name
                - address: The physical address/location of the project
                - description: A text description of the project
                - created_by: Username of the project creator
                - created_at: Formatted creation date and time
                - statistics: Dictionary containing project statistics
        :rtype: dict
        
        :raises ValueError: If no project has been set using set_project()
        
        :example:
            >>> client.set_project("project_123abc")
            >>> info = client.get_project_info()
            >>> print(info["name"])
            'Bridge Inspection'
            >>> print(info["created_at"])
            '2023-04-15 14:30:22'
        """
        if not self.project:
            raise ValueError("Project not set")
        return {
            "id": self.project["id"],
            "name": self.project["profile"]["name"],
            "address": self.project["location"]["address"],
            "description": self.project.get("description", ""),
            "created_by": self.project.get("created_by", ""),
            "created_at": ts2date(self.project["created_at"]),
            "statistics": self.project.get("statistics", {}),
        }

    ################################################################################################
    # CRUD Regions
    ################################################################################################
    def add_region(self, region_name: str):
        """
        Add a new region to the current project.
        
        Regions in T2D2 represent geographical or logical divisions within a project,
        allowing for better organization of assets and data.
        
        :param region_name: The name of the region to add
        :type region_name: str
        
        :return: A dictionary containing the API response with status and region data
        :rtype: dict
        
        :raises ValueError: If no project has been set using set_project()
        
        :example:
            >>> response = client.add_region("North Tower")
            >>> print(response["success"])
            True
            >>> print(response["data"]["name"])
            'North Tower'
        """
        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/categories/regions"
        json_data = self.request(url, RequestType.POST, data={"name": region_name})
        return json_data

    def update_region(self, region_name: str, new_region: dict):
        """
        Update an existing region in the current project.
        
        This method finds a region by name and updates its properties with the provided values.
        
        :param region_name: The name of the region to update
        :type region_name: str
        :param new_region: Dictionary containing the updated region properties
        :type new_region: dict
        
        :return: A dictionary containing the API response with status and updated region data
        :rtype: dict
        
        :raises ValueError: If no project has been set or if the specified region cannot be found
        
        :example:
            >>> updated_data = {"name": "North Tower A", "description": "Northern section of the tower"}
            >>> response = client.update_region("North Tower", updated_data)
            >>> print(response["data"]["name"])
            'North Tower A'
        """
        if not self.project:
            raise ValueError("Project not set")

        region_id = None
        for region in self.project["regions"]:
            if region["name"] == region_name:
                region_id = region["_id"]
                break

        if region_id is None:
            raise ValueError("Region not found")

        url = f"{self.project['id']}/categories/regions/{region_id}"
        json_data = self.request(url, RequestType.PUT, data=new_region)
        return json_data
    
    def get_regions(self):
        """
        Retrieve all regions from the current project.
        
        This method returns a list of all regions defined in the project.
        
        :return: A list of dictionaries, each containing region details
        :rtype: list
        
        :raises ValueError: If no project has been set using set_project()
        
        :example:
            >>> regions = client.get_regions()
            >>> for region in regions:
            ...     print(f"Region: {region['name']}, ID: {region['_id']}")
            Region: North Tower, ID: reg_123abc
            Region: South Tower, ID: reg_456def
        """
        if not self.project:
            raise ValueError("Project not set")
        
        # If regions are already in the project dictionary, return them
        if "regions" in self.project and self.project["regions"]:
            return self.project["regions"]
        
        # Otherwise fetch regions from the API
        url = f"{self.project['id']}/categories/regions"
        json_data = self.request(url, RequestType.GET)
        return json_data["data"]["region_list"]


    ################################################################################################
    # CRUD Assets
    ################################################################################################
    def get_assets(self, asset_type=1, asset_ids=None):
        """
        Retrieve a list of assets based on specified IDs and asset type.
        
        This method fetches detailed information about specific assets in the project.
        
        :param asset_type: The numeric type of assets to retrieve
        :type asset_type: int
        :default asset_type: 1
        
        :param asset_ids: A list of asset IDs to retrieve
        :type asset_ids: list
        :default asset_ids: None
        
        :return: A list of dictionaries, each containing detailed information about an asset
        :rtype: list
        
        :raises ValueError: If no project has been set or the API request fails
        
        :example:
            >>> image_ids = ["img_123", "img_456"]
            >>> images = client.get_assets(asset_type=1, asset_ids=image_ids)
            >>> for img in images:
            ...     print(f"Image: {img['filename']}, URL: {img['url']}")
        
        :note: If asset_ids is None, an empty list is returned
        """
        if asset_ids is None:
            return []

        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/assets"
        payload = {"asset_type": asset_type, "asset_ids": asset_ids}
        json_data = self.request(url, RequestType.POST, data=payload)
        return json_data["data"]

    def add_assets(self, payload):
        """
        Add assets to the current project.
        
        This generic method adds assets of various types to the project using
        a structured payload that defines the assets' properties.
        
        :param payload: A dictionary containing the assets to add
        :type payload: dict
        
        :return: A dictionary containing the API response with status and created asset data
        :rtype: dict
        
        :raises ValueError: If no project has been set or the API request fails
        
        :example:
            >>> asset_data = {
            ...     "project_id": client.project["id"],
            ...     "asset_type": 1,
            ...     "assets": [
            ...         {
            ...             "name": "Bridge Side View",
            ...             "filename": "bridge_side.jpg",
            ...             "url": "path/to/file.jpg",
            ...             "size": {"filesize": 1024000}
            ...         }
            ...     ]
            ... }
            >>> response = client.add_assets(asset_data)
            >>> print(response["success"])
            True
        
        :note: This is a lower-level method. Consider using specialized methods like
            upload_images() or upload_reports() for specific asset types.
        """
        url = f"{self.project['id']}/assets/bulk.create"
        return self.request(url, RequestType.POST, data=payload)

    def download_assets(
        self, asset_ids, asset_type=1, download_dir="./", original_filename=False
    ):
        """
        Download assets from the project to the local filesystem.
        
        This method retrieves assets by their IDs, downloads them from storage,
        and saves them to the specified directory.
        
        :param asset_ids: A list of asset IDs to download
        :type asset_ids: list
        
        :param asset_type: The numeric type of assets to download
        :type asset_type: int
        :default asset_type: 1
        
        :param download_dir: The local directory path where assets will be saved
        :type download_dir: str
        :default download_dir: "./" (current directory)
        
        :param original_filename: Whether to use the original filename (True) or 
                                generate a name based on the asset ID (False)
        :type original_filename: bool
        :default original_filename: False
        
        :return: A dictionary mapping asset IDs to their local file paths
        :rtype: dict
        
        :raises ValueError: If no project has been set, if some assets are not found,
                        or if downloading fails
        
        :example:
            >>> image_ids = ["img_123", "img_456"]
            >>> file_paths = client.download_assets(
            ...     asset_ids=image_ids,
            ...     asset_type=1,
            ...     download_dir="./images",
            ...     original_filename=True
            ... )
            >>> for asset_id, path in file_paths.items():
            ...     print(f"Asset {asset_id} downloaded to {path}")
        """
        if not self.project:
            raise ValueError("Project not set")

        assets = self.get_assets(asset_type, asset_ids)
        if len(assets) != len(asset_ids):
            raise ValueError("Some assets not found")

        output = {}
        for asset in assets:
            url = asset["url"]
            if original_filename:
                file_name = asset["filename"]
            else:
                ext = os.path.splitext(asset["filename"])[1]
                file_name = f"img_{asset['id']}{ext}"
            file_path = os.path.join(download_dir, file_name)
            output[asset["id"]] = file_path
            response = download_file(url, file_path)
            if not response["success"]:
                raise ValueError(response["message"])

        return output

    ################################################################################################
    # CRUD Images
    ################################################################################################
    def upload_images(self, image_paths, image_type=1, params=None):
        """
        Upload images to the current project.
        
        This method uploads images from local paths to the project's S3 storage
        and registers them with the T2D2 API.
        
        :param image_paths: A list of local file paths for the images to upload
        :type image_paths: list
        
        :param image_type: The type of image to upload (1=regular images, 3=orthomosaics)
        :type image_type: int
        :default image_type: 1
        
        :param params: Additional parameters to include in the upload, which may include:
                    - region: Dict with region_id and name (e.g., {"region_id": "id123", "name": "default"})
                    - notes: Text notes for the images
                    - tag_ids: List of tag IDs to associate with the images
                    - scale: Image scale information
        :type params: dict or None
        :default params: None
        
        :return: A dictionary containing the API response with status and created image data
        :rtype: dict
        
        :raises ValueError: If no project has been set or the upload fails
        
        :example:
            >>> image_paths = ["./images/bridge1.jpg", "./images/bridge2.jpg"]
            >>> params = {
            ...     "region": {"region_id": "region123", "name": "North Section"},
            ...     "notes": "Initial inspection images",
            ...     "tag_ids": [1234, 4567]
            ... }
            >>> response = client.upload_images(image_paths, params=params)
            >>> print(response["success"])
            True
        """

        if not self.project:
            logger.error("Cannot upload images: Project not set")
            raise ValueError("Project not set")

        logger.info(f"Uploading {len(image_paths)} images (type: {image_type})")
        
        # Upload images to S3
        assets = []
        image_root = "images"
        if image_type == 3:
            image_root = "orthomosaics"

        for idx, file_path in enumerate(image_paths, 1):
            logger.debug(f"Uploading image {idx}/{len(image_paths)}: {file_path}")
            base, ext = os.path.splitext(os.path.basename(file_path))
            filename = f"{base}_{random_string(6)}{ext}"
            s3_path = (
                self.s3_base_url
                + f"/projects/{self.project['id']}/{image_root}/{filename}"
            )
            result = upload_file(file_path, s3_path)
            if result.get("success", False):
                logger.debug(f"Successfully uploaded to S3: {filename}")
                assets.append(
                    {
                        "name": base,
                        "filename": base + ext,
                        "url": filename,
                        "size": {"filesize": os.path.getsize(file_path)},
                    }
                )
            else:
                logger.warning(f"Failed to upload {file_path} to S3: {result.get('message', 'Unknown error')}")

        logger.info(f"Successfully uploaded {len(assets)}/{len(image_paths)} images to S3")
        
        # Add images to project
        logger.debug("Registering images with T2D2 API")
        payload = {
            "project_id": self.project["id"],
            "asset_type": 1,
            "image_type": image_type,
            "assets": assets,
        }

        if params is not None:
            payload.update(params)

        result = self.add_assets(payload)
        logger.info(f"Successfully registered {len(assets)} images with T2D2 API")
        return result

    def get_images(self, image_ids=None, params=None):
        """
        Retrieve images from the current project.
        
        This method fetches image data either for specific image IDs or for all images in the project
        if no IDs are specified.
        
        :param image_ids: A list of specific image IDs to retrieve, or None to get all images
        :type image_ids: list or None
        :default image_ids: None
        
        :param params: Additional parameters to filter the image results, which may include:
                    - region_id: Filter images by region
                    - tag_ids: Filter images by associated tags
                    - date_range: Filter images by capture date
                    - limit/offset: Pagination parameters
                    - search: Search term for filtering images
                    - sortBy: Sort order (e.g. "id:-1" for descending)
                    - page: Page number for pagination
                    - regions: List of region IDs to filter by
                    - image_types: List of image type IDs to filter by
        :type params: dict or None
        :default params: None
        
        :return: A list of dictionaries, each containing detailed information about an image
        :rtype: list
        
        :raises ValueError: If no project has been set or the request fails
        
        :example:
            >>> # Get all images in a project
            >>> all_images = client.get_images()
            >>> print(f"Total images: {len(all_images)}")
            
            >>> # Get specific images by ID
            >>> specific_images = client.get_images(["img_123", "img_456"])
            >>> for img in specific_images:
            ...     print(f"Image: {img['filename']}, Region: {img['region']['name']}")
            
            >>> # Get images with filtering
            >>> filtered_images = client.get_images(params={
            ...     "search": "",
            ...     "sortBy": "id:-1",
            ...     "page": 1,
            ...     "limit": 10,
            ...     "regions": ["6841697bf8ab7deb15ededc8"],
            ...     "image_types": [1, 2, 4]
            ... })
            >>> print(f"Images in region: {len(filtered_images)}")
        """
        if not self.project:
            logger.error("Cannot get images: Project not set")
            raise ValueError("Project not set yet")

        # all images in project
        if image_ids is None:
            logger.info("Fetching all images from project")
            base_url = f"{self.project['id']}/images"
            if params:
                logger.debug(f"Using filters: {params}")
                # Convert list values to JSON strings
                formatted_params = {}
                for k, v in params.items():
                    if isinstance(v, list):
                        formatted_params[k] = json.dumps(v)
                    else:
                        formatted_params[k] = v
                query_string = "&".join([f"{k}={v}" for k, v in formatted_params.items()])
                url = f"{base_url}?{query_string}"
            else:
                url = base_url
            json_data = self.request(url, RequestType.GET)
            results = json_data["data"]["image_list"]
            logger.info(f"Retrieved {len(results)} images")
            return results

        # Specified image_ids
        logger.info(f"Fetching {len(image_ids)} specific images")
        results = []
        for img_id in image_ids:
            logger.debug(f"Fetching image: {img_id}")
            url = f"{self.project['id']}/images/{img_id}"
            json_data = self.request(url, RequestType.GET)
            results.append(json_data["data"])
        logger.info(f"Retrieved {len(results)} images")
        return results

    def update_images(self, image_ids, payload):
        """
        Update properties of existing images in the current project.
        
        This method allows you to modify image metadata such as region, tags, and notes.
        
        :param image_ids: A list of image IDs to update
        :type image_ids: list
        
        :param payload: A dictionary containing the updates to apply, which may include:
                    - region_id: ID of the region to associate with the images
                    - name: Updated name for the images
                    - tag_ids: List of tag IDs to associate with the images
                    - notes: Updated text notes for the images
                    - scale: Updated image scale information
        :type payload: dict
        
        :return: A dictionary containing the API response with status and updated image data
        :rtype: dict
        
        :raises ValueError: If no project has been set or the update fails
        
        :example:
            >>> image_ids = ["img_123", "img_456"]
            >>> updates = {
            ...     "region_id": "region789",
            ...     "name": "Updated Image Name",
            ...     "notes": "These images show deterioration in the north section",
            ...     "tag_ids": [123, 4567]
            ... }
            >>> response = client.update_images(image_ids, updates)
            >>> print(response["success"])
            True
    """
        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/images/bulk.update"
        payload["image_ids"] = image_ids
        payload["project_id"] = self.project["id"]
        return self.request(url, RequestType.PUT, data=payload)

    def delete_images(self, image_ids):
        """
        Delete images from the current project.
        
        This method removes images from the project based on their IDs. This operation
        is irreversible and will also delete all associated annotations and data.
        
        :param image_ids: A list of image IDs to delete
        :type image_ids: list
        
        :return: A dictionary containing the API response with status and result information
        :rtype: dict
        
        :raises ValueError: If no project has been set or the delete operation fails
        
        :example:
            >>> image_ids = ["img_123", "img_456"]
            >>> response = client.delete_images(image_ids)
            >>> print(response["success"])
            True
            >>> print(f"Deleted {len(response['data']['deleted_ids'])} images")
        
        :warning: This operation cannot be undone. All image data and associated annotations
                will be permanently removed from the project.
        """
        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/images/bulk.delete"
        payload = {"image_ids": image_ids}
        return self.request(url, RequestType.DELETE, data=payload)

    ################################################################################################
    # CRUD Drawings
    ################################################################################################
    def upload_drawings(self, drawing_paths, params=None):
        """
        Upload drawings to the current project.
        
        This method uploads drawing files from local paths to the project's S3 storage
        and registers them with the T2D2 API.
        
        :param drawing_paths: A list of local file paths for the drawings to upload
        :type drawing_paths: list
        :param params: Additional parameters to include in the upload, which may include:
                    - name: Name for the drawings
                    - tag_ids: List of tag IDs to associate with the drawings
                    - region: Dict with region_id and name (e.g., {"region_id": "6281d46733c74ef808a4812e", "name": "default"})
        :type params: dict or None
        :default params: None
        
        :return: A dictionary containing the API response with status and created drawing data
        :rtype: dict
        
        :raises ValueError: If no project has been set or the upload fails
        
        :example:
            >>> drawing_paths = ["./drawings/floor_plan.pdf", "./drawings/elevation.dwg"]
            >>> params = {
            ...     "name": "Building Plans",
            ...     "tag_ids": [123, 456],
            ...     "region": {"region_id": "6281d46733c74ef808a4812e", "name": "default"}
            ... }
            >>> response = client.upload_drawings(drawing_paths, params=params)
            >>> print(response["success"])
            True
        """
        if not self.project:
            raise ValueError("Project not set")

        # Upload drawings to S3
        assets = []
        for file_path in drawing_paths:
            base, ext = os.path.splitext(os.path.basename(file_path))
            filename = f"{base}_{random_string(6)}{ext}"
            s3_path = (
                self.s3_base_url + f"/projects/{self.project['id']}/drawings/{filename}"
            )
            upload_file(file_path, s3_path)
            assets.append(
                {
                    "name": base,
                    "filename": base + ext,
                    "url": filename,
                    "size": {"filesize": os.path.getsize(file_path)},
                }
            )

        # Add drawings to project
        url = f"{self.project['id']}/assets/bulk.create"
        payload = {"project_id": self.project["id"], "asset_type": 2, "assets": assets}
        
        if params is not None:
            payload.update(params)
            
        res = self.request(url, RequestType.POST, data=payload)

        return res

    def get_drawings(self, drawing_ids=None, params=None):
        """
        Retrieve drawings from the current project.
        
        This method fetches drawing data either for specific drawing IDs or for all drawings in the project
        if no IDs are specified.
        
        :param drawing_ids: A list of specific drawing IDs to retrieve, or None to get all drawings
        :type drawing_ids: list or None
        :default drawing_ids: None
        
        :param params: Additional parameters to filter the drawing results, which may include:
                    - region_id: Filter drawings by region
                    - tag_ids: Filter drawings by associated tags
                    - date_range: Filter drawings by upload date
                    - limit/offset: Pagination parameters
        :type params: dict or None
        :default params: None
        
        :return: A list of dictionaries, each containing detailed information about a drawing
        :rtype: list
        
        :raises ValueError: If no project has been set or the request fails
        
        :example:
            >>> # Get all drawings in a project
            >>> all_drawings = client.get_drawings()
            >>> print(f"Total drawings: {len(all_drawings)}")
            
            >>> # Get specific drawings by ID
            >>> specific_drawings = client.get_drawings(["dwg_123", "dwg_456"])
            >>> for dwg in specific_drawings:
            ...     print(f"Drawing: {dwg['filename']}, Region: {dwg['region']['name']}")
            
            >>> # Get drawings with filtering
            >>> filtered_drawings = client.get_drawings(params={"region_id": "region123"})
            >>> print(f"Drawings in region: {len(filtered_drawings)}")
        """
        if not self.project:
            raise ValueError("Project not set")

        # all drawings in project
        if drawing_ids is None:
            url = f"{self.project['id']}/drawings"
            json_data = self.request(url, RequestType.GET, params=params)
            results = json_data["data"]["drawing_list"]
            return results

        results = []
        for dwg_id in drawing_ids:
            url = f"{self.project['id']}/drawings/{dwg_id}"
            json_data = self.request(url, RequestType.GET, params=params)
            results.append(json_data["data"])
        return results

    def update_drawings(self, drawing_ids, payload):
        """
        Update properties of existing drawings in the current project.
        
        This method allows you to modify drawing metadata such as region, tags, and name.
        
        :param drawing_ids: A list of drawing IDs to update
        :type drawing_ids: list
        
        :param payload: A dictionary containing the updates to apply, which may include:
                    - region_id: ID of the region to associate with the drawings
                    - name: Updated name for the drawings
                    - tag_ids: List of tag IDs to associate with the drawings
                    - region: Dict with region_id and name (e.g., {"region_id": "6281d46733c74ef808a4812e", "name": "default"})
        :type payload: dict
        
        :return: A dictionary containing the API response with status and updated drawing data
        :rtype: dict
        
        :raises ValueError: If no project has been set or the update fails
        
        :example:
            >>> drawing_ids = ["dwg_123", "dwg_456"]
            >>> updates = {
            ...     "name": "Updated Drawing Name",
            ...     "tag_ids": [123, 456],
            ...     "region": {"region_id": "6281d46733c74ef808a4812e", "name": "default"}
            ... }
            >>> response = client.update_drawings(drawing_ids, updates)
            >>> print(response["success"])
            True
        """
        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/drawings/bulk.update"
        payload["drawing_ids"] = drawing_ids
        payload["project_id"] = self.project["id"]
        return self.request(url, RequestType.PUT, data=payload)

    def delete_drawings(self, drawing_ids):
        """
        Delete drawings from the current project.
        
        This method removes drawings from the project based on their IDs. This operation
        is irreversible and will also delete all associated data.
        
        :param drawing_ids: A list of drawing IDs to delete
        :type drawing_ids: list
        
        :return: A dictionary containing the API response with status and result information
        :rtype: dict
        
        :raises ValueError: If no project has been set or the delete operation fails
        
        :example:
            >>> drawing_ids = ["dwg_123", "dwg_456"]
            >>> response = client.delete_drawings(drawing_ids)
            >>> print(response["success"])
            True
            >>> print(f"Deleted {len(response['data']['deleted_ids'])} drawings")
        
        :warning: This operation cannot be undone. All drawing data and associated
                information will be permanently removed from the project.
        """
        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/drawings/bulk.delete"
        payload = {"drawing_ids": drawing_ids}
        return self.request(url, RequestType.DELETE, data=payload)

    ################################################################################################
    # CRUD Videos
    ################################################################################################
    def upload_videos(self, video_paths):
        """Upload videos to the current project.
        
        This method uploads a list of videos to the current project. Each video is first
        uploaded to S3 storage and then registered as an asset in the project.
        
        :param video_paths: A list of file paths to the videos to be uploaded
        :type video_paths: list of str
        
        :return: Response from the API containing information about the uploaded videos
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        
        :example:
        
        >>> video_paths = ['/path/to/video1.mp4', '/path/to/video2.mp4']
        >>> response = client.upload_videos(video_paths)
        >>> print(response)
        """

        if not self.project:
            raise ValueError("Project not set")

        # Upload images to S3
        assets = []
        for file_path in video_paths:
            base, ext = os.path.splitext(os.path.basename(file_path))
            filename = f"{base}_{random_string(6)}{ext}"
            s3_path = (
                self.s3_base_url + f"/projects/{self.project['id']}/drawings/{filename}"
            )
            upload_file(file_path, s3_path)
            assets.append(
                {
                    "name": base,
                    "filename": base + ext,
                    "url": filename,
                    "size": {"filesize": os.path.getsize(file_path)},
                }
            )

        # Add images to project
        url = f"{self.project['id']}/assets/bulk.create"
        payload = {"project_id": self.project["id"], "asset_type": 4, "assets": assets}
        res = self.request(url, RequestType.POST, data=payload)

        return res

    def get_videos(self, video_ids=None, params=None):
        """Retrieve videos from the current project.
        
        This method queries the T2D2 API to fetch details for videos in the current project.
        If video_ids is provided, it returns details for those specific videos.
        If video_ids is None, all videos in the project are returned.
        
        :param video_ids: A list of video IDs for which details are to be retrieved, defaults to None
        :type video_ids: list of str, optional
        :param params: Additional query parameters to pass to the API, defaults to None
        :type params: dict, optional
        
        :return: A list of dictionaries, each containing details of a video
        :rtype: list of dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> # Get all videos in project
        >>> all_videos = client.get_videos()
        >>> 
        >>> # Get specific videos by IDs
        >>> video_ids = ['abc123', 'def456']
        >>> specific_videos = client.get_videos(video_ids)
        >>> print(specific_videos)
        [{'id': 'abc123', 'title': 'Video One', ...}, {'id': 'def456', ...}]
        """
        if not self.project:
            raise ValueError("Project not set")

        # all videos in project
        if video_ids is None:
            url = f"{self.project['id']}/videos"
            json_data = self.request(url, RequestType.GET, params=params)
            results = json_data["data"]["video_list"]
            return results

        results = []
        for video_id in video_ids:
            url = f"{self.project['id']}/videos/{video_id}"
            json_data = self.request(url, RequestType.GET, params=params)
            results.append(json_data["data"])
        return results

    def update_videos(self, video_ids, payload):
        """Update multiple videos in the current project.
        
        This method updates the details of videos identified by their IDs in the current project.
        The updates are specified in the payload dictionary and are applied to each video.
        
        :param video_ids: A list of video IDs to update
        :type video_ids: list of str
        :param payload: Dictionary containing the fields to update and their new values
        :type payload: dict
        
        :return: Response from the API containing information about the updated videos
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> video_ids = ['abc123', 'def456']
        >>> payload = {'title': 'New Title', 'description': 'Updated description'}
        >>> response = client.update_videos(video_ids, payload)
        >>> print(response)
        """
        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/videos/bulk.update"
        payload["video_ids"] = video_ids
        payload["project_id"] = self.project["id"]
        return self.request(url, RequestType.PUT, data=payload)

    def delete_videos(self, video_ids):
        """
            Delete videos from the current project.
            
            This method removes videos identified by their IDs from the current project.
            This operation is irreversible.
            
            :param video_ids: A list of video IDs to delete from the project
            :type video_ids: list of str
            
            :return: Response from the API indicating which videos were successfully deleted
            :rtype: dict
            
            :raises ValueError: If no project is currently set or if video_ids is empty or not a list
            :raises ConnectionError: If there is a problem connecting to the T2D2 API
            
            :example:
            
            >>> video_ids = ['abc123', 'def456']
            >>> response = client.delete_videos(video_ids)
            >>> print(response)
        """
        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/videos/bulk.delete"
        payload = {"video_ids": video_ids}
        return self.request(url, RequestType.DELETE, data=payload)

    ################################################################################################
    # CRUD 3D
    ################################################################################################
    def upload_threed(self, threed_paths):
        """
        Upload 3D model files to the current project.
        
        This method takes a list of file paths pointing to 3D model files and uploads them
        to the T2D2 project. Each file is uploaded to a storage location and then registered
        within the project with a unique identifier.
        
        :param threed_paths: A list of file paths to 3D model files
        :type threed_paths: list of str
        
        :return: Response from the API containing information about the uploaded models
        :rtype: dict
        
        :raises FileNotFoundError: If any path in threed_paths does not point to an existing file
        :raises ConnectionError: If there is a problem connecting to the API or storage location
        :raises ValueError: If threed_paths is empty, not a list, or if no project is set
        
        :example:
        
        >>> threed_paths = ['/path/to/model1.obj', '/path/to/model2.stl']
        >>> response = client.upload_threed(threed_paths)
        >>> print(response)
        """

        if not self.project:
            raise ValueError("Project not set")

        # Upload images to S3
        assets = []
        for file_path in threed_paths:
            base, ext = os.path.splitext(os.path.basename(file_path))
            filename = f"{base}_{random_string(6)}{ext}"
            s3_path = (
                self.s3_base_url
                + f"/projects/{self.project['id']}/3d_models/{filename}"
            )
            upload_file(file_path, s3_path)
            assets.append(
                {
                    "name": base,
                    "filename": base + ext,
                    "url": filename,
                    "size": {"filesize": os.path.getsize(file_path)},
                }
            )

        # Add images to project
        url = f"{self.project['id']}/assets/bulk.create"
        payload = {"project_id": self.project["id"], "asset_type": 6, "assets": assets}
        res = self.request(url, RequestType.POST, data=payload)

        return res

    def get_threed(self, model_ids=None, params=None):
        """
        Retrieve 3D models from the current project.
        
        This method queries the project for details on 3D models. If model_ids is provided,
        it returns details for those specific models. If model_ids is None, all 3D models
        in the project are returned.
        
        :param model_ids: A list of 3D model IDs to retrieve, defaults to None
        :type model_ids: list of str, optional
        :param params: Additional query parameters to pass to the API, defaults to None
        :type params: dict, optional
        
        :return: A list of dictionaries, each containing details of a 3D model
        :rtype: list of dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> # Get all 3D models in project
        >>> all_models = client.get_threed()
        >>> 
        >>> # Get specific models by IDs
        >>> model_ids = ['model1_id', 'model2_id']
        >>> specific_models = client.get_threed(model_ids)
        >>> print(specific_models)
        [{'id': 'model1_id', 'name': 'model1.obj', 'size': '2MB'}, {'id': 'model2_id', ...}]
        """

        if not self.project:
            raise ValueError("Project not set")

        # all videos in project
        if model_ids is None:
            url = f"{self.project['id']}/3d-models"
            json_data = self.request(url, RequestType.GET, params=params)
            results = json_data["data"]["model_list"]
            return results

        results = []
        for model_id in model_ids:
            url = f"{self.project['id']}/3d-models/{model_id}"
            json_data = self.request(url, RequestType.GET, params=params)
            results.append(json_data["data"])
        return results

    def update_threed(self, model_ids, payload):
        """
        Update 3D models in the current project.
        
        This method updates the details of 3D models identified by their IDs in the current project.
        The updates are specified in the payload dictionary and are applied to each model.
        
        :param model_ids: A list of 3D model IDs to update
        :type model_ids: list of str
        :param payload: Dictionary containing the fields to update and their new values
        :type payload: dict
        
        :return: Response from the API containing information about the updated models
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> model_ids = ['model1_id', 'model2_id']
        >>> payload = {'name': 'New Model Name', 'description': 'Updated description'}
        >>> response = client.update_threed(model_ids, payload)
        >>> print(response)
        """

        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/3d-models/bulk.update"
        payload["model_ids"] = model_ids
        payload["project_id"] = self.project["id"]
        return self.request(url, RequestType.PUT, data=payload)

    def delete_threed(self, model_ids):
        """
        Delete 3D models from the current project.
        
        This method removes 3D models identified by their IDs from the current project.
        This operation is irreversible.
        
        :param model_ids: A list of 3D model IDs to delete from the project
        :type model_ids: list of str
        
        :return: Response from the API indicating which models were successfully deleted
        :rtype: dict
        
        :raises ValueError: If no project is currently set or if model_ids is empty or not a list
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> model_ids = ['model1_id', 'model2_id']
        >>> response = client.delete_threed(model_ids)
        >>> print(response)
        """

        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/3d-models/bulk.delete"
        payload = {"model_ids": model_ids}
        return self.request(url, RequestType.DELETE, data=payload)

    ################################################################################################
    # CRUD Reports
    ################################################################################################
    def upload_reports(self, report_paths):
        """
        Upload report files to the current project.
        
        This method takes a list of file paths pointing to report files (PDF or docx)
        and uploads them to the T2D2 project. Each file is uploaded to a storage location
        and then registered within the project.
        
        :param report_paths: A list of file paths to report files (PDF or docx)
        :type report_paths: list of str
        
        :return: Response from the API containing information about the uploaded reports
        :rtype: dict
        
        :raises ValueError: If report_paths is empty, not a list, or if no project is set
        :raises FileNotFoundError: If any path does not point to an existing file
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> report_path = "./reports/report_123.pdf"
        >>> response = client.upload_reports([report_path])
        >>> print(response)
        """

        if not self.project:
            raise ValueError("Project not set")

        # Upload images to S3
        assets = []
        for file_path in report_paths:
            base, ext = os.path.splitext(os.path.basename(file_path))
            filename = f"{base}_{random_string(6)}{ext}"
            s3_path = (
                self.s3_base_url + f"/projects/{self.project['id']}/reports/{filename}"
            )
            response = upload_file(file_path, s3_path)
            if not response["success"]:
                raise ValueError(response["message"])

            assets.append(
                {
                    "name": base,
                    "filename": base + ext,
                    "url": filename,
                    "size": {"filesize": os.path.getsize(file_path)},
                }
            )

        # Add images to project
        url = f"{self.project['id']}/assets/bulk.create"
        payload = {"project_id": self.project["id"], "asset_type": 5, "assets": assets}
        res = self.request(url, RequestType.POST, data=payload)

        return res

    def get_reports(self, report_ids=None, params=None):
        """
        Retrieve reports from the current project.
        
        This method queries the project for details on reports identified by their
        unique IDs. It returns information about each report, including metadata.
        
        :param report_ids: A list of report IDs for which details are to be retrieved, defaults to None
        :type report_ids: list of str, optional
        :param params: Additional query parameters to pass to the API, defaults to None
        :type params: dict, optional
        
        :return: A list of dictionaries, each containing details of a report
        :rtype: list of dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> report_ids = ['report1_id', 'report2_id']
        >>> report_details = client.get_reports(report_ids)
        >>> print(report_details)
        [{'id': 'report1_id', 'title': 'Report One', 'creation_date': '2023-01-01'}, ...]
        """
        if report_ids is None:
            return []

        if not self.project:
            raise ValueError("Project not set")

        results = []
        for report_id in report_ids:
            url = f"{self.project['id']}/reports/{report_id}"
            json_data = self.request(url, RequestType.GET, params=params)
            results.append(json_data["data"])
        return results

    def update_reports(self, report_ids, payload):
        """
        Update reports in the current project.
            
        This method updates the details of reports identified by their IDs in the current project.
        The updates are specified in the payload dictionary and are applied to each report.
        
        :param report_ids: A list of report IDs to update
        :type report_ids: list of str
        :param payload: Dictionary containing the fields to update and their new values
        :type payload: dict
        
        :return: Response from the API containing information about the updated reports
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> report_ids = ['report1_id', 'report2_id']
        >>> payload = {'title': 'Updated Report Title', 'status': 'Completed'}
        >>> response = client.update_reports(report_ids, payload)
        >>> print(response)
        """
        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/reports/bulk.update"
        payload["report_ids"] = report_ids
        payload["project_id"] = self.project["id"]
        return self.request(url, RequestType.PUT, data=payload)

    def delete_reports(self, report_ids):
        """
        Delete reports from the current project.
        
        This method removes reports identified by their IDs from the current project.
        This operation is irreversible.
        
        :param report_ids: A list of report IDs to delete from the project
        :type report_ids: list of str
        
        :return: Response from the API indicating which reports were successfully deleted
        :rtype: dict
        
        :raises ValueError: If no project is currently set or if report_ids is empty or not a list
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> report_ids = ['report1_id', 'report2_id']
        >>> response = client.delete_reports(report_ids)
        >>> print(response)
        """
        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/reports/bulk.delete"
        payload = {"report_ids": report_ids}
        return self.request(url, RequestType.DELETE, data=payload)

    def generate_condition_report_document(self, image_ids=None, output_path='condition_report.docx', padding_percent=0.2):
        """
        Generate a Word document with condition report for images and their annotations.
        
        Creates a Word document with:
        
        * First page: Project details and T2D2 information
        * Subsequent pages: One page per annotation showing:
        
          * Cropped annotated image (top)
          * Original image (middle)
          * Metadata table (bottom) with file name, link, condition, region, tags
        
        :param image_ids: List of specific image IDs to include, or None to include all images
        :type image_ids: list or None
        :default image_ids: None
        
        :param output_path: Path where the Word document should be saved
        :type output_path: str
        :default output_path: 'condition_report.docx'
        
        :param padding_percent: Percentage of padding to add around cropped annotations (0.0 to 1.0)
        :type padding_percent: float
        :default padding_percent: 0.2
        
        :return: Path to the generated document
        :rtype: str
        
        :raises ValueError: If no project is currently set
        
        :example:
            >>> # Generate report for all images
            >>> doc_path = client.generate_condition_report_document()
            >>> print(f"Report saved to: {doc_path}")
            
            >>> # Generate report for specific images
            >>> doc_path = client.generate_condition_report_document(
            ...     image_ids=['img_123', 'img_456'],
            ...     output_path='custom_report.docx'
            ... )
        """
        if not self.project:
            logger.error("Cannot generate condition report: Project not set")
            raise ValueError("Project not set")
        
        logger.info(f"Starting condition report generation. Output: {output_path}, Image IDs: {image_ids or 'All'}")
        
        # Get project info
        project_info = self.get_project_info()
        logger.debug(f"Project: {project_info['name']} (ID: {project_info['id']})")
        
        # Get images
        if image_ids is None:
            logger.info("Fetching all images for condition report")
            images = self.get_images()
        else:
            logger.info(f"Fetching {len(image_ids)} specific images for condition report")
            images = self.get_images(image_ids=image_ids)
        
        if not images:
            logger.error("No images found for condition report")
            raise ValueError("No images found")
        
        logger.info(f"Found {len(images)} images to process")
        
        # Prepare image data for ImageAnnotationCropper
        image_data_list = []
        for img in images:
            image_id = img.get('id')
            
            # Get annotations for this image
            try:
                annotations = self.get_annotations(image_id=image_id)
                logger.debug(f"Image {image_id}: Found {len(annotations)} annotations")
            except Exception as e:
                logger.warning(f"Failed to get annotations for image {image_id}: {e}")
                annotations = []
            
            # Filter visible annotations only
            visible_annotations = [ann for ann in annotations if ann.get('visible', True)]
            
            if not visible_annotations:
                continue
            
            # Get image URL - try different possible fields
            image_url = img.get('url') or img.get('s3_url') or img.get('image_url') or img.get('file_url')
            if not image_url:
                # Try to construct from s3_base_url if available
                if hasattr(self, 's3_base_url') and img.get('filename'):
                    image_url = f"{self.s3_base_url}/{img['filename']}"
                else:
                    logger.warning(f"No URL found for image {image_id}, skipping...")
                    continue
            
            # Get image dimensions
            width = img.get('width') or img.get('image_width') or 1920
            height = img.get('height') or img.get('image_height') or 1080
            
            # Format image data for ImageAnnotationCropper
            image_data = {
                'url': image_url,
                'info': {
                    'width': width,
                    'height': height
                },
                'annotations': visible_annotations,
                'image_id': img.get('id'),
                'filename': img.get('filename', ''),
                'region': img.get('region', {}),
                'tags': img.get('tags', [])
            }
            image_data_list.append(image_data)
        
        if not image_data_list:
            logger.error("No images with annotations found for condition report")
            raise ValueError("No images with annotations found")
        
        logger.info(f"Processing {len(image_data_list)} images with annotations")
        
        # Initialize ImageAnnotationCropper
        logger.debug("Initializing ImageAnnotationCropper")
        cropper = ImageAnnotationCropper(image_data_list)
        logger.info("Downloading images...")
        cropper.download_images()
        
        # Create Word document
        logger.info("Creating Word document structure")
        doc = Document()
        
        # Set document margins - optimized for better layout
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # First page: Title, Logo, then Project details
        # Add title first with better spacing
        title = doc.add_heading('Condition Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(12)
        title.paragraph_format.space_after = Pt(18)
        
        # Add T2D2 logo after title
        logo_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 't2d2_image', 't2d2.png')
        if os.path.exists(logo_path):
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_para.paragraph_format.space_before = Pt(6)
            logo_para.paragraph_format.space_after = Pt(18)
            run = logo_para.add_run()
            run.add_picture(logo_path, width=Inches(2))
        
        # Add project information in a better formatted way
        # Calculate number of rows needed
        num_rows = 4  # Project, Project ID, Project Link, Address (always present)
        if project_info.get('description'):
            num_rows += 1
        num_rows += 1  # Created (always present)
        if project_info.get('created_by'):
            num_rows += 1
        
        project_table = doc.add_table(rows=num_rows, cols=2)
        project_table.style = 'Light Grid Accent 1'
        
        # Set column widths - optimized for better layout
        project_table.columns[0].width = Inches(1.5)
        project_table.columns[1].width = Inches(4.5)
        
        # Fill project information table
        row_idx = 0
        project_table.cell(row_idx, 0).text = 'Project:'
        project_table.cell(row_idx, 1).text = project_info['name']
        row_idx += 1
        
        project_table.cell(row_idx, 0).text = 'Project ID:'
        project_table.cell(row_idx, 1).text = str(project_info['id'])
        row_idx += 1
        
        # Add project link
        project_table.cell(row_idx, 0).text = 'Project Link:'
        project_link_url = f"{self.base_url.replace('/api/', '')}/project/{self.project['id']}"
        link_paragraph = project_table.cell(row_idx, 1).paragraphs[0]
        link_paragraph.clear()
        add_hyperlink(link_paragraph, project_link_url, project_link_url)
        row_idx += 1
        
        project_table.cell(row_idx, 0).text = 'Address:'
        project_table.cell(row_idx, 1).text = project_info['address']
        row_idx += 1
        
        if project_info.get('description'):
            project_table.cell(row_idx, 0).text = 'Description:'
            project_table.cell(row_idx, 1).text = project_info['description']
            row_idx += 1
        
        project_table.cell(row_idx, 0).text = 'Created:'
        project_table.cell(row_idx, 1).text = project_info['created_at']
        row_idx += 1
        
        if project_info.get('created_by'):
            project_table.cell(row_idx, 0).text = 'Created By:'
            project_table.cell(row_idx, 1).text = project_info['created_by']
        
        # Format project table cells with better spacing and alignment
        for row in project_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style.font.size = Pt(11)
                    paragraph.paragraph_format.space_before = Pt(6)
                    paragraph.paragraph_format.space_after = Pt(6)
                    if cell == row.cells[0]:  # First column (labels)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.bold = True
                    else:  # Second column (values)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add page break after project details
        doc.add_page_break()
        
        # Process each image and annotation
        figure_counter = 1
        first_annotation = True
        
        for img_idx, (image_data, pil_image) in enumerate(cropper.images):
            if pil_image is None:
                continue
            
            image_width = image_data['info']['width']
            image_height = image_data['info']['height']
            annotations = image_data['annotations']
            image_id = image_data.get('image_id', '')
            filename = image_data.get('filename', '')
            region = image_data.get('region', {})
            tags = image_data.get('tags', [])
            
            # Get portal URL (construct from base_url and project/image IDs)
            portal_url = f"{self.base_url.replace('/api/', '')}/project/{self.project['id']}/images/{image_id}"
            
            # Process each annotation
            for ann in annotations:
                # Add page break before each annotation to ensure one condition per page
                # (skip for first annotation since we already have a page break after project details)
                if not first_annotation:
                    doc.add_page_break()
                first_annotation = False
                # Crop annotation
                cropped_img, crop_bbox = cropper.crop_annotation(
                    pil_image, ann, image_width, image_height, padding_percent
                )
                
                if cropped_img is None or crop_bbox is None:
                    continue
                
                # Draw annotation on cropped image
                cropped_with_annotation = cropper.draw_annotation_on_image(
                    cropped_img, ann, image_width, image_height, crop_bbox
                )
                
                # Add minimal spacing at top of page (optimized for one page per condition)
                doc.add_paragraph()
                
                # Add cropped annotated image (top) - optimized size to fit on one page
                cropped_para = doc.add_paragraph()
                cropped_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cropped_para.paragraph_format.space_before = Pt(6)
                cropped_para.paragraph_format.space_after = Pt(3)
                
                # Convert PIL image to bytes for docx
                cropped_bytes = BytesIO()
                cropped_with_annotation.save(cropped_bytes, format='PNG')
                cropped_bytes.seek(0)
                
                # Add image to document (optimized width to fit on one page with other content)
                run = cropped_para.add_run()
                run.add_picture(cropped_bytes, width=Inches(5.5))
                
                # Add figure caption below cropped image with minimal spacing
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.paragraph_format.space_before = Pt(3)
                caption_para.paragraph_format.space_after = Pt(6)
                caption_run = caption_para.add_run(f"Figure {figure_counter}: Annotated Image (Cropped)")
                caption_run.font.size = Pt(9)
                caption_run.italic = True
                figure_counter += 1
                
                # Add minimal spacing between images
                doc.add_paragraph()
                
                # Add original image (middle) - optimized size to fit on one page
                original_para = doc.add_paragraph()
                original_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                original_para.paragraph_format.space_before = Pt(3)
                original_para.paragraph_format.space_after = Pt(3)
                
                # Draw all annotations on original for context
                original_with_annotations = pil_image.copy()
                for visible_ann in annotations:
                    original_with_annotations = cropper.draw_annotation_on_image(
                        original_with_annotations, visible_ann, image_width, image_height
                    )
                
                # Convert original image to bytes
                original_bytes = BytesIO()
                original_with_annotations.save(original_bytes, format='PNG')
                original_bytes.seek(0)
                
                # Add original image (optimized width to fit on one page)
                run = original_para.add_run()
                run.add_picture(original_bytes, width=Inches(3.5))
                
                # Add figure caption below original image with minimal spacing
                caption2_para = doc.add_paragraph()
                caption2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption2_para.paragraph_format.space_before = Pt(3)
                caption2_para.paragraph_format.space_after = Pt(6)
                caption2_run = caption2_para.add_run(f"Figure {figure_counter}: Original Image with All Annotations")
                caption2_run.font.size = Pt(9)
                caption2_run.italic = True
                figure_counter += 1
                
                # Add minimal spacing before table
                doc.add_paragraph()
                
                # Create metadata table
                table = doc.add_table(rows=5, cols=2)
                table.style = 'Light Grid Accent 1'
                
                # Set column widths for better layout
                table.columns[0].width = Inches(1.8)
                table.columns[1].width = Inches(4.2)
                
                # Get annotation details
                annotation_class = ann.get('annotation_class', {})
                class_name = annotation_class.get('annotation_class_long_name', 'N/A')
                condition = ann.get('condition', {})
                condition_name = condition.get('rating_name', 'N/A') if condition else 'N/A'
                region_name = region.get('name', 'N/A') if region else 'N/A'
                tags_str = ', '.join([tag.get('name', '') for tag in tags if isinstance(tag, dict)]) or 'N/A'
                if not tags_str or tags_str == '':
                    tags_str = 'N/A'
                
                # Fill table
                table.cell(0, 0).text = 'File Name'
                table.cell(0, 1).text = filename or 'N/A'
                
                # Add clickable hyperlink for portal URL
                table.cell(1, 0).text = 'Link to Photo in T2D2 Portal'
                link_cell = table.cell(1, 1)
                link_paragraph = link_cell.paragraphs[0]
                link_paragraph.clear()
                add_hyperlink(link_paragraph, portal_url, portal_url)
                
                table.cell(2, 0).text = 'Condition'
                table.cell(2, 1).text = condition_name
                
                table.cell(3, 0).text = 'Region'
                table.cell(3, 1).text = region_name
                
                table.cell(4, 0).text = 'Tags'
                table.cell(4, 1).text = tags_str
                
                # Format table cells with optimized styling for one page layout
                for row_idx, row in enumerate(table.rows):
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.style.font.size = Pt(9)
                            # Reduced padding to fit on one page
                            paragraph.paragraph_format.space_before = Pt(4)
                            paragraph.paragraph_format.space_after = Pt(4)
                            # Make first column (labels) bold and align left
                            if cell == row.cells[0]:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                for run in paragraph.runs:
                                    run.bold = True
                            else:  # Second column (values)
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # No page break needed here - next annotation will add its own page break at the start
        
        # Save document
        logger.info(f"Saving condition report document to: {output_path}")
        doc.save(output_path)
        logger.info(f"Condition report document saved successfully: {output_path}")
        
        return output_path

    ################################################################################################
    # CRUD Tags
    ################################################################################################
    def get_tags(self, params=None):
        """Retrieve tags from the current project.
        
        This method queries the project for all tags or specific tags based on the
        provided parameters. It returns information about each tag, including its
        name and associated metadata.
        
        :param params: Additional query parameters to filter tags, defaults to None
        :type params: dict, optional
        
        :return: Tag data from the API response
        :rtype: list of dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> # Get all tags
        >>> all_tags = client.get_tags()
        >>> print(all_tags)
        >>> 
        >>> # Get tags with specific filter
        >>> filtered_tags = client.get_tags(params={'name': 'example'})
        >>> print(filtered_tags)
        """
        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/tags"
        json_data = self.request(url, RequestType.GET, params=params)
        return json_data["data"]

    def add_tags(self, tags):
        """Add new tags to the current project.
        
        This method creates new tags within the project. Each tag is defined by
        a name. If a tag already exists, a warning is printed and the operation
        for that tag is skipped.
        
        :param tags: Either a single tag name (string) or a list of tag names
        :type tags: str or list of str
        
        :return: List of API responses for each successfully created tag
        :rtype: list of dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> # Add a single tag
        >>> result = client.add_tags('New Tag')
        >>> print(result)
        >>> 
        >>> # Add multiple tags
        >>> result = client.add_tags(['Tag 1', 'Tag 2', 'Tag 3'])
        >>> print(result)
        """
        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/tags"
        if isinstance(tags, str):
            tags = [tags]

        results = []
        for tag in tags:
            payload = {"name": tag}
            try:
                # Use POST to create new tag
                result = self.request(url, RequestType.POST, data=payload)
                results.append(result)
            except Exception as e:
                logger.warning(f"Tag already exists: {tag} - {e}")

        return results

    # def delete_tags(self, tag_ids):
    #     """
    #     Delete tags from the current project.
        
    #     This method removes tags identified by their IDs from the current project.
    #     This operation is irreversible.
        
    #     :param tag_ids: A list of tag IDs to delete
    #     :type tag_ids: list of int
        
    #     :return: Response from the API indicating the status of the deletion operation
    #     :rtype: dict
        
    #     :raises ValueError: If no project is currently set
    #     :raises ConnectionError: If there is a problem connecting to the T2D2 API
    #     """
    #     if not self.project:
    #         raise ValueError("Project not set")

    #     # Make sure tag_ids are properly formatted
    #     validated_tag_ids = []
    #     for tag_id in tag_ids:
    #         # If it's already a number, add it directly
    #         if isinstance(tag_id, int) or (isinstance(tag_id, str) and tag_id.isdigit()):
    #             validated_tag_ids.append(int(tag_id))
    #         else:
    #             # If it's a MongoDB ObjectId string, try to find the corresponding numeric ID
    #             # from the tags list (this would require having the tags list cached or fetching it)
    #             pass  # Implement if needed

    #     url = f"{self.project['id']}/tags/bulk.delete"
    #     payload = {"ids": validated_tag_ids}
    #     result = self.request(url, RequestType.DELETE, data=payload)
        
    #     return result

    # def update_tags(self, tag_updates):
    #     """
    #     Update existing tags in the current project.
        
    #     This method updates the properties of existing tags identified by their IDs.
    #     Each update specification must include the tag ID and the new property values.
        
    #     :param tag_updates: A list of dictionaries, each containing a tag ID and the properties to update
    #     :type tag_updates: list of dict
        
    #     :return: Response from the API containing information about the updated tags
    #     :rtype: dict
        
    #     :raises ValueError: If no project is currently set
    #     :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
    #     :example:
        
    #     >>> tag_updates = [
    #     ...     {'id': 'tag1', 'name': 'Updated Tag 1'},
    #     ...     {'id': 'tag2', 'name': 'Updated Tag 2'}
    #     ... ]
    #     >>> result = client.update_tags(tag_updates)
    #     >>> print(result)
    #     """
    #     if not self.project:
    #         raise ValueError("Project not set")
        
    #     url = f"{self.project['id']}/tags/bulk.update"
    #     payload = {"updates": tag_updates}
    #     result = self.request(url, RequestType.PUT, data=payload)
        
    #     return result
    
    ################################################################################################
    # CRUD Annotation Classes
    ################################################################################################
    def get_materials(self):
        """Retrieve materials from the current project.
        
        This method queries the project's database for a list of all materials that 
        have been defined within the project.
        
        :return: Information about materials in the project
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> materials = client.get_materials()
        >>> print(materials)
        {'material_list': ['Concrete', 'Steel', 'Wood']}
        """
        if not self.project:
            raise ValueError("Project not set")

        url = "material"
        params = {"project_id": self.project["id"]}
        json_data = self.request(url, RequestType.GET, params=params)
        return json_data["data"]

    def get_annotation_classes(self, params=None):
        """Retrieve annotation classes from the current project.
        
        This method queries the project's database for annotation classes defined within 
        the project. It returns information about each class including ID, name, and attributes.
        
        :param params: Additional query parameters to filter classes, defaults to None
        :type params: dict, optional
        
        :return: Information about annotation classes in the project
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> annotation_classes = client.get_annotation_classes()
        >>> print(annotation_classes)
        {'label_list': [{'id': 'class1', 'name': 'Vehicle', 'attributes': ['color', 'make']}, ...]}
        """
        if not self.project:
            raise ValueError("Project not set")

        url = "annotation-class"
        params_ = {
            "project_id": self.project["id"],
            "scope": "DEFAULT",
            "sortBy": "id:asc",
        }
        if params is not None:
            params_.update(params)

        print(params_)

        json_data = self.request(url, RequestType.GET, params=params_)
        return json_data["data"]

    def add_annotation_class(self, name, color=None, materials=None):
        """Add a new annotation class to the current project.
        
        This method creates a new annotation class with the specified name, color,
        and associated materials.
        
        :param name: The name of the annotation class to create
        :type name: str
        :param color: The color for the annotation class, defaults to a random color
        :type color: str, optional
        :param materials: A list of materials associated with this class, defaults to an empty list
        :type materials: list of str, optional
        
        :return: Response from the API containing information about the created class
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> # Create a basic annotation class
        >>> result = client.add_annotation_class('Pedestrian')
        >>> print(result)
        >>> 
        >>> # Create an annotation class with specific color and materials
        >>> result = client.add_annotation_class('Crack', '#FF0000', ['Concrete', 'Asphalt'])
        >>> print(result)
        """
        if not self.project:
            raise ValueError("Project not set")

        if materials is None:
            materials = []

        if color is None:
            color = random_color()

        url = "annotation-class/create-annotation-class"
        payload = {
            "project_id": self.project["id"],
            "name": name,
            "materials": materials,
            "color": color,
        }
        results = self.request(url, RequestType.POST, data=payload)


    def delete_annotation_classes(self, annotation_class_ids):
        """Delete annotation classes from the current project.
        
        This method removes annotation classes identified by their IDs from the project.
        This operation is irreversible.
        
        :param annotation_class_ids: One or more annotation class IDs to delete
        :type annotation_class_ids: list of str or int
        
        :return: Response from the API indicating the status of the deletion operation
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> # Delete a single annotation class
        >>> result = client.delete_annotation_classes('class1_id')
        >>> print(result)
        >>> 
        >>> # Delete multiple annotation classes
        >>> result = client.delete_annotation_classes(['class1_id', 'class2_id'])
        >>> print(result)
        """
        if not self.project:
            raise ValueError("Project not set")

        if isinstance(annotation_class_ids, int):
            annotation_class_ids = [annotation_class_ids]

        if len(annotation_class_ids) == 0:
            return {"success": False, "message": "No annotation class ids provided"}

        url = f"{self.project['id']}/annotation-class/bulk.delete"
        payload = {"annotation_class_ids": annotation_class_ids}
        results = self.request(url, RequestType.DELETE, data=payload)

        return results

    # TODO: Update Annotation Classes
    ################################################################################################
    # CRUD Annotations
    ################################################################################################
    def get_annotations(self, image_id=None, params=None):
        """
        Retrieve annotations from the current project.
        
        This method queries the project for annotations. If image_id is provided, 
        it retrieves annotations for that specific image. If image_id is None,
        it retrieves annotations for all images matching the provided params.
        
        :param image_id: The ID of the image for which to retrieve annotations, defaults to None
        :type image_id: str, optional
        :param params: Additional query parameters to filter images, defaults to None
        :type params: dict, optional
        
        :return: A list of dictionaries, each containing details of an annotation
        :rtype: list of dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> # Get annotations for a specific image
        >>> image_id = 'image123'
        >>> annotations = client.get_annotations(image_id)
        >>> print(annotations)
        >>> 
        >>> # Get annotations for images matching certain criteria
        >>> params = {'region_id': 'region1', 'date_from': '2023-01-01'}
        >>> all_annotations = client.get_annotations(params=params)
        >>> print(all_annotations)
        """

        if not self.project:
            logger.error("Cannot get annotations: Project not set")
            raise ValueError("Project not set")

        if params is None:
            params = {}

        if image_id is None:
            logger.info("Fetching annotations for all images matching params")
            images = self.get_images(params=params)
            image_ids = [img["id"] for img in images]
            logger.debug(f"Found {len(image_ids)} images to fetch annotations from")
        else:
            logger.info(f"Fetching annotations for image: {image_id}")
            image_ids = [image_id]

        images = self.get_images(image_ids=image_ids, params=params)
        annotations = []
        for img in images:
            img_annotations = img.get("annotations", [])
            annotations.extend(img_annotations)
            logger.debug(f"Image {img.get('id')}: Found {len(img_annotations)} annotations")

        logger.info(f"Retrieved {len(annotations)} total annotations")
        return annotations

    def delete_annotations(self, image_id, annotation_ids=None):
        """
        Delete annotations from an image in the current project.
        
        This method removes annotations from a specific image. If annotation_ids is provided,
        it deletes those specific annotations. If annotation_ids is None, it deletes all
        annotations in the image.
        
        :param image_id: The ID of the image from which to delete annotations
        :type image_id: str
        :param annotation_ids: A list of annotation IDs to delete, defaults to None (all annotations)
        :type annotation_ids: list of str, optional
        
        :return: Response from the API indicating the status of the deletion operation
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> # Delete specific annotations from an image
        >>> image_id = 'image123'
        >>> annotation_ids = ['ann1', 'ann2']
        >>> result = client.delete_annotations(image_id, annotation_ids)
        >>> print(result)
        >>> 
        >>> # Delete all annotations from an image
        >>> image_id = 'image123'
        >>> result = client.delete_annotations(image_id)
        >>> print(result)
        """

        if not self.project:
            logger.error("Cannot delete annotations: Project not set")
            raise ValueError("Project not set")

        if annotation_ids is None:
            # Get annotation_ids for all annotations in image
            logger.info(f"Deleting all annotations from image: {image_id}")
            annotations = self.get_annotations(image_id)
            annotation_ids = [ann["id"] for ann in annotations]
            logger.debug(f"Found {len(annotation_ids)} annotations to delete")
        else:
            logger.info(f"Deleting {len(annotation_ids)} specific annotations from image: {image_id}")

        payload = {
            "project_id": self.project["id"],
            "image_id": image_id,
            "annotation_ids": annotation_ids,
        }

        result = self.request("annotation", RequestType.DELETE, data=payload)
        logger.info(f"Successfully deleted annotations from image: {image_id}")
        return result

    def add_annotations(self, image_id, annotations):
        """
        Add new annotations to an image in the current project.
        
        This method creates new annotations for a specific image. Each annotation
        is defined by its properties such as class, geometry, and attributes.
        
        :param image_id: The ID of the image to which annotations will be added
        :type image_id: str
        :param annotations: A list of annotation definitions to be added
        :type annotations: list of dict
        
        :return: Response from the API containing information about the created annotations
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> image_id = 'image123'
        >>> annotations_to_add = [
        ...     {
        ...         'annotation_class_id': 'class1',
        ...         'coordinates': [[100, 100], [200, 100], [200, 200], [100, 200]],
        ...         'attributes': {'severity': 'high'}
        ...     },
        ...     {
        ...         'annotation_class_id': 'class2',
        ...         'coordinates': [[300, 300], [400, 400]],
        ...     }
        ... ]
        >>> result = client.add_annotations(image_id, annotations_to_add)
        >>> print(result)
        """
        if not self.project:
            logger.error("Cannot add annotations: Project not set")
            raise ValueError("Project not set")

        logger.info(f"Adding {len(annotations)} annotations to image: {image_id}")
        url = "annotation"
        payload = {
            "project_id": self.project["id"],
            "image_id": image_id,
            "annotations": annotations,
        }

        results = self.request(url, RequestType.POST, data=payload)
        logger.info(f"Successfully added annotations to image: {image_id}")

        return results

    # TODO: Update Annotations
    ################################################################################################
    # CRUD Geotags
    ################################################################################################
    def get_geotags(self, drawing_id, params=None):
        """Retrieve geotags associated with a drawing in the current project.
        
        This method queries the project's database for geotags associated with the specified
        drawing. Additional filter parameters can be provided through the params argument.
        
        :param drawing_id: The ID of the drawing for which to retrieve geotags
        :type drawing_id: str
        :param params: Additional query parameters to filter geotags, defaults to None
        :type params: dict, optional
        
        :return: A list of dictionaries, each containing details of a geotag
        :rtype: list of dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> drawing_id = 'drawing123'
        >>> geotags = client.get_geotags(drawing_id)
        >>> print(geotags)
        [{'id': 'geotag1', 'latitude': 40.7138, 'longitude': -74.0065, 'tags': ['tag1']}, ...]
        """
        if drawing_id is None:
            return []

        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/geotags?drawing_id={drawing_id}"
        json_data = self.request(url, RequestType.GET, params=params)

        return json_data["data"]

    def add_geotags(self, drawing_id, geotags):
        """Add new geotags to a drawing in the current project.
        
        This method creates new geotags associated with the specified drawing.
        Each geotag is defined by its coordinates and can include additional
        metadata.
        
        :param drawing_id: The ID of the drawing to which the geotags will be added
        :type drawing_id: str
        :param geotags: A list of geotag definitions to be added
        :type geotags: list of dict
        
        :return: Response from the API containing information about the created geotags
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> drawing_id = 'drawing123'
        >>> geotags_to_add = [
        ...     {'latitude': 40.7128, 'longitude': -74.0060, 'tags': ['tag1', 'tag2']},
        ...     {'latitude': 34.0522, 'longitude': -118.2437}
        ... ]
        >>> result = client.add_geotags(drawing_id, geotags_to_add)
        >>> print(result)
        """
        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/geotags/bulk.create"
        payload = {
            "drawing_id": drawing_id,
            "geotags": geotags,
        }
        results = self.request(url, RequestType.POST, data=payload)

        return results

    def delete_geotags(self, drawing_id, geotag_ids):
        """Delete specified geotags from a drawing in the current project.
        
        This method removes geotags identified by their IDs from a specified drawing.
        This operation is irreversible.
        
        :param drawing_id: The ID of the drawing from which the geotags will be deleted
        :type drawing_id: str
        :param geotag_ids: A list of geotag IDs to delete
        :type geotag_ids: list of str
        
        :return: Response from the API indicating the status of the deletion operation
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> drawing_id = 'drawing123'
        >>> geotag_ids = ['geotag1', 'geotag2']
        >>> result = client.delete_geotags(drawing_id, geotag_ids)
        >>> print(result)
        """
        if not self.project:
            raise ValueError("Project not set")

        url = f"{self.project['id']}/geotags/bulk.delete"
        payload = {
            "drawing_id": drawing_id,
            "geotag_ids": geotag_ids,
        }
        results = self.request(url, RequestType.POST, data=payload)

        return results

    ################################################################################################
    # CRUD Downloads
    ################################################################################################
    def upload_downloads(self, file_paths):
        """
        Upload files to the downloads section of the current project.
        
        This method uploads files to the downloads section of the current project.
        Each file is uploaded to a pre-specified storage location and is associated
        with the project.
        
        :param file_paths: A list of file paths to be uploaded
        :type file_paths: list of str
        
        :return: Dictionary indicating success or failure of the upload operation
        :rtype: dict
        
        :raises ValueError: If no project is currently set or if upload fails
        :raises FileNotFoundError: If any path in file_paths does not point to an existing file
        :raises ConnectionError: If there is a problem connecting to the storage service
        
        :example:
        
        >>> file_paths = ['/path/to/report.pdf', '/path/to/data.zip']
        >>> result = client.upload_downloads(file_paths)
        >>> print(result)
        {'success': True, 'message': 'Files uploaded'}
        """
        
        if not self.project:
            raise ValueError("Project not set")

        for file_path in file_paths:
            filename = os.path.basename(file_path)
            s3_path = (
                self.s3_base_url
                + f"/projects/{self.project['id']}/downloads/{filename}"
            )
            response = upload_file(file_path, s3_path)
            if not response["success"]:
                raise ValueError(response["message"])

        return {"success": True, "message": "Files uploaded"}

    ################################################################################################
    # Classes and Conditions methods
    ################################################################################################
    def get_classes(self):
        """
        Retrieve annotation classes with associated conditions.
        
        This method retrieves all annotation classes from the current project and
        maps them to their corresponding conditions. It returns a list of condition
        objects enriched with the name of the annotation class they belong to.
        
        :return: A list of dictionaries, each containing condition details with associated class name
        :rtype: list of dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> classes = client.get_classes()
        >>> print(classes)
        [{'id': 'cond1', 'name': 'Crack', 'annotation_class_id': 'class1', ...}, ...]
        """
        if not self.project:
            raise ValueError("Project not set")

        # Get full class data
        class_map = {}
        classes = self.get_annotation_classes()
        for lbl in classes["label_list"]:
            class_map[lbl["id"]] = lbl

        url = f"{self.project['id']}/conditions"
        json_data = self.request(url, RequestType.GET)
        output = []
        for lbl in json_data["data"]["condition_list"]:
            lbl["name"] = class_map[lbl["annotation_class_id"]]["name"]
            output.append(lbl)
        return output  # Fixed missing return statement

    def summarize_images(self, params=None):
        """
        Summarize images in the current project by region, date, and tags.
        
        This method analyzes all images in the project and groups them by region,
        capture date, and assigned tags. It returns counts for each group category.
        
        :param params: Additional query parameters to filter images, defaults to None
        :type params: dict, optional
        
        :return: A dictionary with three subdictionaries: 'region_group', 'date_group', and 'tag_group'
                Each containing counts of images per category
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> summary = client.summarize_images()
        >>> print(summary)
        {
            'region_group': {'Region A': 10, 'Region B': 5},
            'date_group': {'2023-01-15': 7, '2023-01-16': 8},
            'tag_group': {'Priority': 3, 'Review': 12}
        }
        """
        if not self.project:
            raise ValueError("Project not set")

        # Return image summary
        images = self.get_images(params=params)
        if len(images) == 0:
            return {
                "region_group": {},
                "date_group": {},
                "tag_group": {},
            }
        # Group by region, date and tags
        region_group, date_group, tag_group = (
            defaultdict(int),
            defaultdict(int),
            defaultdict(int),
        )
        for img in images:
            img_region = img["region"]["name"]
            img_date = ts2date(img["captured_date"]).split(" ")[0]
            img_tags = img["tags"]

            region_group[img_region] += 1
            date_group[img_date] += 1
            for img_tag in img_tags:
                tag_group[img_tag["name"]] += 1

        return {
            "region_group": region_group,
            "date_group": date_group,
            "tag_group": tag_group,
        }

    def summarize_conditions(self, params=None):
        """
        Summarize annotations and their conditions by region, label, and rating.
        
        This method analyzes all annotations in the project, grouped by region, and summarizes
        them based on annotation class label and condition rating. For each combination,
        it calculates total count, cumulative length, total area, and collects annotation IDs.
        
        :param params: Additional query parameters to filter images, defaults to None
        :type params: dict, optional
        
        :return: A nested dictionary with regions as top-level keys, each containing summaries
                for label-rating combinations with count, length, area, and annotation_ids
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> summary = client.summarize_conditions()
        >>> print(summary)
        {
            'Region A': {
                ('Crack', 'Severe'): {
                    'count': 5,
                    'length': 120.5,
                    'area': 45.2,
                    'annotation_ids': ['ann1', 'ann2', ...]
                },
                ...
            },
            'Region B': {...}
        }
        """
        if not self.project:
            raise ValueError("Project not set")

        imgs = self.get_images(params=params)

        anns = defaultdict(list)
        for img in imgs:
            reg = img["region"]["name"]
            anns_img = self.get_annotations(image_id=img["id"])
            anns[reg] += anns_img

        result = {}
        for reg, annotations in anns.items():
            sublist = {}
            for ann in annotations:
                label = ann["annotation_class"]["annotation_class_name"]
                condition = ann.get("condition", {})
                if isinstance(condition, dict):
                    rating = condition.get("rating_name", "default")
                else:
                    rating = "default"
                area = ann["area"]
                length = ann["length"]
                ann_id = ann["id"]
                key = (label, rating)
                if key in sublist:
                    sublist[key]["count"] += 1
                    sublist[key]["length"] += length
                    sublist[key]["area"] += area
                    sublist[key]["annotation_ids"].append(ann_id)
                else:
                    sublist[key] = {}
                    sublist[key]["count"] = 1
                    sublist[key]["length"] = length
                    sublist[key]["area"] = area
                    sublist[key]["annotation_ids"] = [ann_id]  # Fixed missing initial annotation ID

            result[reg] = sublist

        return result


    ################################################################################################
    # Notifications
    ################################################################################################

    def slack_notification(self, payload):
        """
        Send a notification to Slack.
    
        This method sends a notification to a Slack channel using the T2D2 API.
        The notification content and destination are specified in the payload.
        
        :param payload: Dictionary containing the notification details
        :type payload: dict
        
        :return: Response from the API indicating whether the notification was sent successfully
        :rtype: dict
        
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> payload = {
        ...     "channel": "#project-updates",
        ...     "message": "New 3D model uploaded",
        ...     "attachments": [{"title": "Model Details", "text": "Model ID: abc123"}]
        ... }
        >>> response = client.slack_notification(payload)
        >>> print(response)
        """


        url = "notifications/slack"
        return self.request(url, RequestType.POST, data=payload)
    

    ################################################################################################
    # AI Models
    ################################################################################################

    def get_ai_models(self):
        """
        Retrieve all AI models available in the current project.
        
        This method fetches all AI models associated with the current project from the T2D2 API.
        It returns a list of dictionaries containing details about each AI model.
        
        :return: A list of dictionaries, each containing AI model details
        :rtype: list of dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> models = client.get_ai_models()
        >>> print(models)
        """
        if not self.project:
            raise ValueError("Project not set")
        
        url = f"{self.project['id']}/ai-models"
        response = self.request(url, RequestType.GET)
        
        # Return just the data list if response has the expected structure
        if isinstance(response, dict) and 'data' in response:
            return response['data']
        return response
    
    def get_ai_model_by_id(self, model_id):
        """
        Retrieve details of a specific AI model by its ID.
        
        This method fetches detailed information about a specific AI model from the T2D2 API.
        It returns a dictionary containing all the information available for the specified AI model.
        
        :param model_id: The ID of the AI model to retrieve
        :type model_id: str
        
        :return: A dictionary containing AI model details
        :rtype: dict
        
        """

        if not self.project:
            raise ValueError("Project not set")
        
        url = f"{self.project['id']}/ai-models/{model_id}"
        return self.request(url, RequestType.GET)

    def create_ai_model(self, name, config, labels, shape, docker=None, project_ids=None):
        """
        Create a new AI model in the current project.
        
        This method creates a new AI model with the specified configuration, labels, and shape.
        The model will be associated with the current project.
        
        :param name: The name of the AI model to create
        :type name: str
        
        :param config: Configuration dictionary containing model paths and settings.
                      Should include 'weights_path', 'classes', and 'config' fields.
                      The API expects config.paths.config to be present.
        :type config: dict
        
        :param labels: List of labels/classes that the model can detect
        :type labels: list of str
        
        :param shape: The shape parameter for the model
        :type shape: int
        
        :param docker: Docker configuration dictionary containing registry, image, and tag.
                      If not provided, defaults to docker.io/t2d2ai/detectron2_inferencer:latest
        :type docker: dict, optional
        
        :param project_ids: List of project IDs to associate with the model. If not provided
                           and a project is set, the current project ID will be added to the list.
                           If not provided and no project is set, an empty list will be used.
        :type project_ids: list, optional
        
        :return: A dictionary containing the created AI model details
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> config = {
        ...     "weights_path": "/models/roof-discoloration/best_model.pth",
        ...     "classes": "./models/roof-discoloration/model_classes.json",
        ...     "config": "/models/roof-discoloration/model_config.json"
        ... }
        >>> labels = ["roof_discoloration"]
        >>> docker_config = {
        ...     "registry": "docker.io",
        ...     "image": "t2d2ai/detectron2_inferencer",
        ...     "tag": "latest"
        ... }
        >>> model = client.create_ai_model("Roof discoloration", config, labels, 4, docker=docker_config)
        >>> print(f"Created model: {model['name']} with ID: {model['id']}")
        """
        if not self.project:
            raise ValueError("Project not set")
        
        # Ensure config has the required structure for the API
        formatted_config = {
            "weights_path": config.get("weights_path", ""),
            "classes": config.get("classes", ""),
            "config": config.get("config", ""),
            "paths": {
                "config": config.get("config", ""),
                "classes": config.get("classes", ""),
                "weights_path": config.get("weights_path", "")
            }
        }
        
        # Set default docker configuration if not provided
        if docker is None:
            docker = {
                "registry": "docker.io",
                "image": "t2d2ai/detectron2_inferencer",
                "tag": "latest"
            }
        
        # Handle project_ids
        if project_ids is None:
            if self.project:
                project_ids = [self.project["id"]]
            else:
                project_ids = []
        
        url = f"{self.project['id']}/ai-models"
        payload = {
            "name": name,
            "config": formatted_config,
            "labels": labels,
            "shape": shape,
            "docker": docker,
            "project_ids": project_ids
        }
        return self.request(url, RequestType.POST, data=payload)

    def update_ai_model(self, model_id, name=None, config=None, labels=None, shape=None):
        """
        Update an existing AI model in the current project.
        
        This method updates the properties of an existing AI model. Only the provided
        parameters will be updated, others will remain unchanged.
        
        :param model_id: The ID of the AI model to update
        :type model_id: int
        
        :param name: Updated name for the AI model
        :type name: str, optional
        
        :param config: Updated configuration dictionary
        :type config: dict, optional
        
        :param labels: Updated list of labels/classes
        :type labels: list of str, optional
        
        :param shape: Updated shape parameter
        :type shape: int, optional
        
        :return: A dictionary containing the updated AI model details
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> # Update just the name
        >>> updated_model = client.update_ai_model(1, name="Updated Model Name")
        >>> print(f"Updated model: {updated_model['name']}")
        
        >>> # Update multiple properties
        >>> updated_model = client.update_ai_model(
        ...     1, 
        ...     name="New Name", 
        ...     labels=["new_label1", "new_label2"]
        ... )
        """
        if not self.project:
            raise ValueError("Project not set")
        
        url = f"{self.project['id']}/ai-models/{model_id}"
        payload = {}
        
        if name is not None:
            payload["name"] = name
        if config is not None:
            # Format config with required paths structure for the API
            formatted_config = {
                "weights_path": config.get("weights_path", ""),
                "classes": config.get("classes", ""),
                "config": config.get("config", ""),
                "paths": {
                    "config": config.get("config", ""),
                    "classes": config.get("classes", ""),
                    "weights_path": config.get("weights_path", "")
                }
            }
            payload["config"] = formatted_config
        if labels is not None:
            payload["labels"] = labels
        if shape is not None:
            payload["shape"] = shape
            
        return self.request(url, RequestType.PUT, data=payload)

    def delete_ai_model(self, model_id):
        """
        Delete one or more AI models from the current project.
        
        This method removes AI model(s) identified by their ID(s) from the current project.
        This operation is irreversible.
        
        :param model_id: The ID of the AI model to delete, or a list of model IDs to delete
        :type model_id: int or list of int
        
        :return: A dictionary containing the API response with status and message
        :rtype: dict
        
        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> # Delete a single model
        >>> response = client.delete_ai_model(1)
        >>> print(response["message"])
        'AI model deleted successfully'
        
        >>> # Delete multiple models
        >>> response = client.delete_ai_model([21, 22, 23])
        >>> print(response["message"])
        
        :warning: This operation cannot be undone. All AI model data and associated
                information will be permanently removed from the project.
        """
        if not self.project:
            raise ValueError("Project not set")
        
        # Handle both single ID and list of IDs
        if isinstance(model_id, list):
            # Delete multiple models
            results = []
            for mid in model_id:
                url = f"{self.project['id']}/ai-models/{mid}"
                result = self.request(url, RequestType.DELETE)
                results.append(result)
            return results
        else:
            # Delete single model
            url = f"{self.project['id']}/ai-models/{model_id}"
            return self.request(url, RequestType.DELETE)
    
    def run_ai_inferencer(self, image_ids, model_id, confidence_threshold=0.5, 
                         replace_annotations=False, sliding_window=False, 
                         whole_image=True, batch_size=1):
        """
        Run AI inference on specified images using a selected model.
        
        This method initiates an AI inference process on a set of images using the specified
        model. It constructs the payload with model details and user parameters, then sends
        the request to start the inference process.
        
        :param image_ids: List of image IDs to run inference on
        :type image_ids: list of int
        
        :param model_id: ID of the AI model to use for inference
        :type model_id: int
        
        :param confidence_threshold: Minimum confidence score for detections (0.0 to 1.0)
        :type confidence_threshold: float
        :default confidence_threshold: 0.5
        
        :param replace_annotations: Whether to replace existing annotations
        :type replace_annotations: bool
        :default replace_annotations: False
        
        :param sliding_window: Whether to use sliding window approach for inference
        :type sliding_window: bool
        :default sliding_window: False
        
        :param whole_image: Whether to process the whole image at once
        :type whole_image: bool
        :default whole_image: True
        
        :param batch_size: Number of images to process in each batch
        :type batch_size: int
        :default batch_size: 1
        
        :return: Response from the API containing the inference job details
        :rtype: dict
        
        :raises ValueError: If no project is set or if required parameters are invalid
        
        :example:
        
        >>> image_ids = [000000, 000000, 000000]
        >>> result = client.run_ai_inferencer(
        ...     image_ids=image_ids,
        ...     model_id=1,
        ...     confidence_threshold=0.6
        ... )
        >>> print(result)
        """
        if not self.project:
            raise ValueError("Project not set")
            
        # Get model details
        model_details = self.get_ai_model_by_id(model_id)
        if not model_details["success"]:
            raise ValueError(f"Failed to get model details: {model_details.get('message', 'Unknown error')}")
            
        model_data = model_details["data"]
        
        # Construct payload
        payload = {
            "confidence_threshold": confidence_threshold,
            "replace_annotations": replace_annotations,
            "whole_image": whole_image,
            "batch_size": batch_size,
            "ai_model": model_id,
            "description": model_data.get("description",""),
            "image_ids": image_ids,
            "labels": model_data["labels"],
            "sliding_window": sliding_window
        }
        
        # Make request to start inference
        url = f"{self.project['id']}/tools/6/start"
        return self.request(url, RequestType.POST, data=payload)
    
    def get_task_list(self):
        """
        Get the list of tasks for the current project.
        
        This method fetches the list of tasks associated with the current project from the T2D2 API.
        It returns a list of dictionaries containing details about each task.
        
        :return: A list of dictionaries, each containing task details
        :rtype: list of dict

        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> tasks = client.get_task_list()
        >>> print(tasks)
        """
        if not self.project:
            raise ValueError("Project not set")
        
        url = f"{self.project['id']}/task"
        return self.request(url, RequestType.GET)
    
    def get_task_by_id(self, task_id):
        """
        Get the details of a specific task by its ID.
        
        This method fetches detailed information about a specific task from the T2D2 API.
        It returns a dictionary containing all the information available for the specified task.
        
        :param task_id: The ID of the task to retrieve
        :type task_id: str
        
        :return: A dictionary containing task details
        :rtype: dict

        :raises ValueError: If no project is currently set
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> task = client.get_task_by_id("9d9f28a5-df16-4f69-bbc9-4dc718ea204c")
        >>> print(task)
        """
        if not self.project:
            raise ValueError("Project not set")
        
        # First get all tasks
        url = f"{self.project['id']}/task"
        response = self.request(url, RequestType.GET)
        
        # Find the task with matching task_id
        for task in response["data"]["task_list"]:
            if task["task_id"] == task_id:
                return task
                
        raise ValueError(f"Task with ID {task_id} not found")

    ################################################################################################
    # CRUD Datasets
    ################################################################################################
    def get_datasets(self, params=None):
        """
        Retrieve datasets from the T2D2 API.
        
        This method fetches all datasets available to the authenticated user.
        It supports filtering and pagination through the params parameter.
        
        :param params: Additional query parameters to filter datasets, which may include:
                    - search: Search term for filtering datasets by name
                    - sortBy: Sort order (e.g. "id:desc" for descending by ID)
                    - page: Page number for pagination
                    - limit: Number of datasets per page
                    - public: Filter by public/private datasets (true/false)
        :type params: dict or None
        :default params: None
        
        :return: A dictionary containing the dataset list and total count
        :rtype: dict
        
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> # Get all datasets
        >>> datasets = client.get_datasets()
        >>> print(f"Total datasets: {datasets['total_datasets']}")
        >>> for dataset in datasets['dataset_list']:
        ...     print(f"Dataset: {dataset['name']}, ID: {dataset['id']}")
        
        >>> # Get datasets with filtering
        >>> filtered_datasets = client.get_datasets(params={
        ...     "search": "test",
        ...     "sortBy": "id:desc",
        ...     "page": 1,
        ...     "limit": 10,
        ...     "public": False
        ... })
        >>> print(f"Found {len(filtered_datasets['dataset_list'])} datasets")
        """
        url = "datasets"
        json_data = self.request(url, RequestType.GET, params=params)
        return json_data["data"]

    def create_dataset(self, name):
        """
        Create a new dataset.
        
        This method creates a new dataset with the specified name.
        The dataset will be private by default and owned by the authenticated user.
        
        :param name: The name of the dataset to create
        :type name: str
        
        :return: A dictionary containing the created dataset details
        :rtype: dict
        
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> dataset = client.create_dataset("My Training Dataset")
        >>> print(f"Created dataset: {dataset['name']} with ID: {dataset['id']}")
        """
        url = "datasets"
        payload = {"name": name}
        json_data = self.request(url, RequestType.POST, data=payload)
        return json_data["data"]

    def delete_datasets(self, dataset_ids):
        """
        Delete datasets by their IDs.
        
        This method removes datasets identified by their IDs from the T2D2 system.
        This operation is irreversible and will also delete all associated data.
        
        :param dataset_ids: A list of dataset IDs to delete
        :type dataset_ids: list of int
        
        :return: A dictionary containing the API response with status and message
        :rtype: dict
        
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> dataset_ids = [8, 9]
        >>> response = client.delete_datasets(dataset_ids)
        >>> print(response["message"])
        'Dataset is deleted successfuly'
        
        :warning: This operation cannot be undone. All dataset data and associated
                information will be permanently removed.
        """
        url = "datasets/bulk.delete"
        payload = {"dataset_ids": dataset_ids}
        return self.request(url, RequestType.DELETE, data=payload)

    def update_dataset_images(self, dataset_id, action, image_ids):
        """
        Add or remove images from a dataset.
        
        This method allows you to add or remove images from an existing dataset.
        The action parameter determines whether to add or remove the specified images.
        
        :param dataset_id: The ID of the dataset to update
        :type dataset_id: int
        
        :param action: The action to perform - either "add" or "remove"
        :type action: str
        
        :param image_ids: A list of image IDs to add or remove from the dataset
        :type image_ids: list of int
        
        :return: A dictionary containing the updated dataset details
        :rtype: dict
        
        :raises ValueError: If action is not "add" or "remove"
        :raises ConnectionError: If there is a problem connecting to the T2D2 API
        
        :example:
        
        >>> # Add images to dataset
        >>> image_ids = [660068, 660069]
        >>> updated_dataset = client.update_dataset_images(8, "add", image_ids)
        >>> print(f"Dataset now has {len(updated_dataset['image_ids'])} images")
        
        >>> # Remove images from dataset
        >>> updated_dataset = client.update_dataset_images(8, "remove", [660068])
        >>> print(f"Dataset now has {len(updated_dataset['image_ids'])} images")
        """
        if action not in ["add", "remove"]:
            raise ValueError("Action must be either 'add' or 'remove'")
        
        url = f"datasets/{dataset_id}"
        payload = {
            "action": action,
            "image_ids": image_ids
        }
        json_data = self.request(url, RequestType.PUT, data=payload)
        return json_data["data"]
